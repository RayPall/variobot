# variobot.py
"""Streamlit app for Vario bot
--------------------------------
• Posílá název modulu do Make přes webhook.
• Přijímá zpětný JSON string (?payload=) a hned zobrazí finální text.
• Umožní stáhnout vygenerovaný obsah jako DOCX.
"""

import io
import json
import urllib.parse
from typing import Any, Dict, Optional

import requests
import streamlit as st
from docx import Document

# ==============  Nastavení  ==============

WEBHOOK_URL = "https://hook.eu2.make.com/6dobqwk57qdm23w6p09pgvnmrrl9qp72"
PAGE_TITLE = "Vario Bot – Landing‑page generátor"
PAGE_ICON = "📝"

st.set_page_config(page_title=PAGE_TITLE, page_icon=PAGE_ICON, layout="centered")
st.title("📝 Generátor (a příjemce) popisků modulů ERP Vario")

# ============== Pomocné funkce ==============

def extract_text(payload_raw: str) -> Optional[str]:
    """Rozbalí %xx sekvence, parse JSON a vrátí hodnotu `result` nebo `text`."""
    try:
        decoded = urllib.parse.unquote_plus(payload_raw)
        data: Dict[str, Any] = json.loads(decoded)
    except (json.JSONDecodeError, TypeError):
        return None
    return data.get("result") or data.get("text")


def build_docx(data: Dict[str, Any]) -> bytes:
    """Vytvoří DOCX v paměti a vrátí ho jako byty."""
    doc = Document()
    title = data.get("module", "Vario modul")
    doc.add_heading(title, level=1)

    # Pokud klíč "Text" existuje, použij ho jako jediný blok.
    if "Text" in data and isinstance(data["Text"], str):
        doc.add_paragraph(data["Text"])
    else:
        for key, val in data.items():
            if key in ("module",):
                continue
            doc.add_heading(str(key), level=2)
            doc.add_paragraph(str(val))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============== 1) Formulář pro odeslání názvu modulu ==============

st.header("🔸 Odeslat název modulu do Make")
with st.form(key="mod_form", clear_on_submit=True):
    module_name = st.text_input("Název modulu", placeholder="Např. Skladové hospodářství")
    submitted = st.form_submit_button("Odeslat")

if submitted:
    if not module_name.strip():
        st.warning("Prosím, vyplňte název modulu.")
        st.stop()

    payload_out = {"module_name": module_name.strip()}
    try:
        resp = requests.post(WEBHOOK_URL, json=payload_out, timeout=10)
        resp.raise_for_status()
    except requests.exceptions.RequestException as exc:
        st.error(f"❌ Nepodařilo se odeslat: {exc}")
    else:
        st.success("✅ Název modulu byl úspěšně odeslán do Make.")
        st.json(payload_out)

# ============== 2) Automatický příjem & zobrazení textu ==============

# Query param payload
query_params = st.experimental_get_query_params()
qp_payload = query_params.get("payload", [None])[0]

final_text = extract_text(qp_payload) if qp_payload else None

if final_text:
    # Pokud jsme payload našli, zobrazíme a nabídneme DOCX.
    st.divider()
    st.success("✅ Text úspěšně přijat z Make webhooku")

    # Nadpis a Markdown
    st.subheader("📄 Výstupní text")
    st.markdown(final_text, unsafe_allow_html=True)

    # Soubor DOCX k downloadu
    docx_bytes = build_docx({"module": "Výstup z Make", "Text": final_text})
    st.download_button(
        label="💾 Stáhnout DOCX",
        data=docx_bytes,
        file_name="landing_page.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Už nemusíme zobrazovat nic dalšího.
    st.stop()

# ============== 3) Manuální JSON fallback (volitelné) ==============

st.divider()
st.header("🔸 Ruční kontrola / JSON fallback")
json_input = st.text_area(
    "Vlož JSON, nebo přidej do URL ?payload= …:",
    value="",
    height=160,
    placeholder='{"result": "Text landing‑page …"}'
)

if st.button("Zobraz JSON"):
    try:
        data_dict = json.loads(json_input)
    except json.JSONDecodeError as err:
        st.error(f"Neplatný JSON: {err}")
        st.stop()

    st.json(data_dict)
    text_part = data_dict.get("result") or data_dict.get("text")
    if text_part:
        st.markdown(text_part, unsafe_allow_html=True)
        docx_bytes = build_docx({"module": "Ruční JSON", "Text": text_part})
        st.download_button(
            "💾 Stáhnout DOCX",
            docx_bytes,
            file_name="landing_manual.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.info("V JSONu není klíč 'result' ani 'text'.")
